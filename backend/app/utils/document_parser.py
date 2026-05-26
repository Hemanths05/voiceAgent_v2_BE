"""
Document Parser Utilities
Extracts text from various document formats for knowledge base ingestion
"""
import io
from typing import Dict, Any, Optional
from pathlib import Path
import PyPDF2
import pandas as pd
from docx import Document
from app.core.logging_config import get_logger
from app.core.exceptions import DocumentParsingError

logger = get_logger(__name__)


class DocumentParser:
    """
    Handles parsing of various document formats

    Supported formats:
    - PDF (.pdf)
    - Text (.txt)
    - Word (.docx)
    - CSV (.csv)
    """

    @staticmethod
    def parse_pdf(file_data: bytes, filename: str = "document.pdf") -> Dict[str, Any]:
        """
        Extract text from PDF file

        Args:
            file_data: PDF file bytes
            filename: Original filename

        Returns:
            Dict with 'text', 'pages', 'metadata'

        Raises:
            DocumentParsingError: If parsing fails
        """
        try:
            pdf_buffer = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_buffer)

            # Extract metadata
            metadata = {
                "filename": filename,
                "format": "pdf",
                "num_pages": len(pdf_reader.pages),
            }

            # Try to get PDF metadata
            try:
                if pdf_reader.metadata:
                    metadata["title"] = pdf_reader.metadata.get("/Title", "")
                    metadata["author"] = pdf_reader.metadata.get("/Author", "")
                    metadata["subject"] = pdf_reader.metadata.get("/Subject", "")
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata: {str(e)}")

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    continue

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise DocumentParsingError(
                    f"No text could be extracted from PDF: {filename}",
                    {"filename": filename, "format": "pdf"}
                )

            logger.info(
                f"Successfully parsed PDF: {filename} "
                f"({metadata['num_pages']} pages, {len(full_text)} chars)"
            )

            return {
                "text": full_text,
                "pages": metadata["num_pages"],
                "metadata": metadata
            }

        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {str(e)}", exc_info=True)
            raise DocumentParsingError(
                f"Failed to parse PDF: {str(e)}",
                {"filename": filename, "format": "pdf"}
            )

    @staticmethod
    def parse_txt(file_data: bytes, filename: str = "document.txt") -> Dict[str, Any]:
        """
        Extract text from plain text file

        Args:
            file_data: Text file bytes
            filename: Original filename

        Returns:
            Dict with 'text', 'metadata'

        Raises:
            DocumentParsingError: If parsing fails
        """
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_data.decode('utf-8')
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed for {filename}, trying latin-1")
                text = file_data.decode('latin-1')

            if not text.strip():
                raise DocumentParsingError(
                    f"Text file is empty: {filename}",
                    {"filename": filename, "format": "txt"}
                )

            metadata = {
                "filename": filename,
                "format": "txt",
                "num_lines": len(text.splitlines()),
            }

            logger.info(
                f"Successfully parsed TXT: {filename} "
                f"({metadata['num_lines']} lines, {len(text)} chars)"
            )

            return {
                "text": text,
                "metadata": metadata
            }

        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Error parsing TXT {filename}: {str(e)}", exc_info=True)
            raise DocumentParsingError(
                f"Failed to parse text file: {str(e)}",
                {"filename": filename, "format": "txt"}
            )

    @staticmethod
    def parse_docx(file_data: bytes, filename: str = "document.docx") -> Dict[str, Any]:
        """
        Extract text from Word document

        Args:
            file_data: DOCX file bytes
            filename: Original filename

        Returns:
            Dict with 'text', 'metadata'

        Raises:
            DocumentParsingError: If parsing fails
        """
        try:
            docx_buffer = io.BytesIO(file_data)
            document = Document(docx_buffer)

            # Extract text from paragraphs
            paragraphs = []
            for para in document.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            table_texts = []
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)

            # Combine all text
            text_parts = paragraphs
            if table_texts:
                text_parts.append("\n[Tables]\n" + "\n".join(table_texts))

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise DocumentParsingError(
                    f"No text could be extracted from DOCX: {filename}",
                    {"filename": filename, "format": "docx"}
                )

            metadata = {
                "filename": filename,
                "format": "docx",
                "num_paragraphs": len(paragraphs),
                "num_tables": len(document.tables),
            }

            # Try to get core properties
            try:
                core_props = document.core_properties
                metadata["title"] = core_props.title or ""
                metadata["author"] = core_props.author or ""
                metadata["subject"] = core_props.subject or ""
            except Exception as e:
                logger.warning(f"Could not extract DOCX metadata: {str(e)}")

            logger.info(
                f"Successfully parsed DOCX: {filename} "
                f"({metadata['num_paragraphs']} paragraphs, "
                f"{metadata['num_tables']} tables, {len(full_text)} chars)"
            )

            return {
                "text": full_text,
                "metadata": metadata
            }

        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {str(e)}", exc_info=True)
            raise DocumentParsingError(
                f"Failed to parse Word document: {str(e)}",
                {"filename": filename, "format": "docx"}
            )

    @staticmethod
    def parse_csv(file_data: bytes, filename: str = "document.csv") -> Dict[str, Any]:
        """
        Extract text from CSV file

        Args:
            file_data: CSV file bytes
            filename: Original filename

        Returns:
            Dict with 'text', 'metadata'

        Raises:
            DocumentParsingError: If parsing fails
        """
        try:
            csv_buffer = io.BytesIO(file_data)

            # Try to read CSV
            try:
                df = pd.read_csv(csv_buffer, encoding='utf-8')
            except UnicodeDecodeError:
                csv_buffer.seek(0)
                df = pd.read_csv(csv_buffer, encoding='latin-1')

            if df.empty:
                raise DocumentParsingError(
                    f"CSV file is empty: {filename}",
                    {"filename": filename, "format": "csv"}
                )

            # Convert to text representation
            # Format: Column headers + rows as key-value pairs
            text_parts = []

            # Add column information
            text_parts.append(f"Columns: {', '.join(df.columns)}\n")

            # Add each row as structured text
            for idx, row in df.iterrows():
                row_text = f"Row {idx + 1}:\n"
                row_text += "\n".join(
                    f"  {col}: {row[col]}"
                    for col in df.columns
                    if pd.notna(row[col])
                )
                text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)

            metadata = {
                "filename": filename,
                "format": "csv",
                "num_rows": len(df),
                "num_columns": len(df.columns),
                "columns": list(df.columns),
            }

            logger.info(
                f"Successfully parsed CSV: {filename} "
                f"({metadata['num_rows']} rows, "
                f"{metadata['num_columns']} columns, {len(full_text)} chars)"
            )

            return {
                "text": full_text,
                "metadata": metadata
            }

        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Error parsing CSV {filename}: {str(e)}", exc_info=True)
            raise DocumentParsingError(
                f"Failed to parse CSV file: {str(e)}",
                {"filename": filename, "format": "csv"}
            )

    @classmethod
    def parse(
        cls,
        file_data: bytes,
        filename: str,
        file_format: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse document based on file format

        Args:
            file_data: File bytes
            filename: Original filename
            file_format: File format (pdf, txt, docx, csv). If None, inferred from filename

        Returns:
            Dict with 'text', 'metadata', and optional 'pages'

        Raises:
            DocumentParsingError: If format is unsupported or parsing fails
        """
        # Infer format from filename if not provided
        if file_format is None:
            file_format = Path(filename).suffix.lower().lstrip('.')

        if not file_format:
            raise DocumentParsingError(
                "Could not determine file format",
                {"filename": filename}
            )

        # Route to appropriate parser
        parsers = {
            "pdf": cls.parse_pdf,
            "txt": cls.parse_txt,
            "docx": cls.parse_docx,
            "doc": cls.parse_docx,  # Treat .doc as .docx (may fail for old format)
            "csv": cls.parse_csv,
        }

        parser = parsers.get(file_format)
        if parser is None:
            raise DocumentParsingError(
                f"Unsupported file format: {file_format}",
                {"filename": filename, "format": file_format}
            )

        logger.info(f"Parsing document: {filename} (format: {file_format})")
        return parser(file_data, filename)

    @staticmethod
    def validate_document_size(
        file_data: bytes,
        max_size_mb: int = 10
    ) -> None:
        """
        Validate document size

        Args:
            file_data: File bytes
            max_size_mb: Maximum allowed size in MB

        Raises:
            DocumentParsingError: If file is too large
        """
        size_mb = len(file_data) / (1024 * 1024)
        if size_mb > max_size_mb:
            raise DocumentParsingError(
                f"Document too large: {size_mb:.2f}MB (max: {max_size_mb}MB)",
                {"size_mb": size_mb, "max_size_mb": max_size_mb}
            )

    @staticmethod
    def validate_text_length(
        text: str,
        min_length: int = 50,
        max_length: int = 1_000_000
    ) -> None:
        """
        Validate extracted text length

        Args:
            text: Extracted text
            min_length: Minimum required length
            max_length: Maximum allowed length

        Raises:
            DocumentParsingError: If text is too short or too long
        """
        text_length = len(text.strip())

        if text_length < min_length:
            raise DocumentParsingError(
                f"Document text too short: {text_length} chars (min: {min_length})",
                {"text_length": text_length, "min_length": min_length}
            )

        if text_length > max_length:
            raise DocumentParsingError(
                f"Document text too long: {text_length} chars (max: {max_length})",
                {"text_length": text_length, "max_length": max_length}
            )


# Export public class
__all__ = ["DocumentParser"]
